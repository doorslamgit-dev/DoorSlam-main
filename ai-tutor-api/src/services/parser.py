# ai-tutor-api/src/services/parser.py
# Document text extraction — multi-format parsing via Docling with legacy fallback.

import logging
import threading
from dataclasses import dataclass, field
from io import BytesIO

from ..config import settings

logger = logging.getLogger(__name__)

# Extensions that Docling can handle
DOCLING_EXTENSIONS = frozenset({
    ".pdf", ".docx", ".pptx", ".xlsx", ".csv",
    ".html", ".htm", ".md", ".txt", ".tex",
    ".png", ".jpg", ".jpeg", ".tiff", ".bmp",
})


@dataclass
class ParsedDocument:
    """Result of parsing a document into text."""

    text: str
    page_texts: list[str] = field(default_factory=list)
    page_count: int = 0
    metadata: dict = field(default_factory=dict)


# --- Docling singleton (lazy, thread-safe) ---

_converter = None
_converter_lock = threading.Lock()


def _get_converter():
    """Create or return the Docling DocumentConverter singleton.

    Uses double-checked locking so the heavy import + model load
    happens at most once across all threads.
    """
    global _converter
    if _converter is not None:
        return _converter

    with _converter_lock:
        if _converter is not None:
            return _converter

        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            PdfPipelineOptions,
            TableFormerMode,
        )
        from docling.document_converter import DocumentConverter, PdfFormatOption

        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_ocr = settings.docling_do_ocr
        pipeline_options.do_table_structure = True
        if settings.docling_table_mode == "accurate":
            pipeline_options.table_structure_options.mode = TableFormerMode.ACCURATE
        else:
            pipeline_options.table_structure_options.mode = TableFormerMode.FAST

        _converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(
                    pipeline_options=pipeline_options,
                ),
            }
        )
        logger.info("Docling DocumentConverter initialised (ocr=%s, table_mode=%s)",
                     settings.docling_do_ocr, settings.docling_table_mode)
        return _converter


def _reset_converter():
    """Reset the Docling singleton — for test isolation only."""
    global _converter
    with _converter_lock:
        _converter = None


# --- Public API ---

def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Extract text from a document based on its file extension.

    When Docling is enabled and the extension is supported, uses Docling
    for structured Markdown output (preserving tables). Falls back to
    legacy parsers per-file if Docling fails or is disabled.

    Supports via Docling: .pdf .docx .pptx .xlsx .csv .html .htm .md .txt .tex
                          .png .jpg .jpeg .tiff .bmp
    Legacy fallback: .pdf (PyMuPDF), .docx (python-docx), .md/.txt (UTF-8)
    """
    lower = filename.lower()
    ext = _get_extension(lower)

    # Try Docling first if enabled and extension is supported
    if settings.docling_enabled and ext in DOCLING_EXTENSIONS:
        try:
            return _parse_with_docling(file_bytes, filename)
        except Exception as exc:
            if settings.docling_fallback and ext in {".pdf", ".docx", ".md", ".txt"}:
                logger.warning(
                    "Docling failed for %s, falling back to legacy parser: %s",
                    filename, exc,
                )
            else:
                raise ValueError(
                    f"Docling parsing failed for {filename}: {exc}"
                ) from exc

    # Legacy parsers
    if ext == ".pdf":
        return _parse_pdf_legacy(file_bytes)
    elif ext == ".docx":
        return _parse_docx_legacy(file_bytes)
    elif ext in {".md", ".txt"}:
        return _parse_text_legacy(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


# --- Docling parser ---

def _parse_with_docling(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse a document using Docling. Returns Markdown-formatted text."""
    from docling.datamodel.base_models import DocumentStream

    converter = _get_converter()
    source = DocumentStream(name=filename, stream=BytesIO(file_bytes))
    result = converter.convert(source)

    # Full document as Markdown (tables preserved as | col | col | syntax)
    full_text = result.document.export_to_markdown()

    # Per-page texts
    page_texts = []
    page_numbers = sorted(result.document.pages.keys())
    for page_num in page_numbers:
        page_md = result.document.export_to_markdown(page_no=page_num)
        page_texts.append(page_md)

    ext = _get_extension(filename.lower())

    return ParsedDocument(
        text=full_text,
        page_texts=page_texts if page_texts else [full_text],
        page_count=len(page_texts) if page_texts else 1,
        metadata={"parser": "docling", "format": ext.lstrip(".")},
    )


# --- Legacy parsers ---

def _parse_pdf_legacy(file_bytes: bytes) -> ParsedDocument:
    """Extract text from PDF using PyMuPDF (plain text, no table support)."""
    import pymupdf

    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    page_texts = []
    for page in doc:
        page_texts.append(page.get_text())

    full_text = "\n\n".join(page_texts)
    page_count = len(page_texts)
    doc.close()

    return ParsedDocument(
        text=full_text,
        page_texts=page_texts,
        page_count=page_count,
        metadata={"parser": "legacy", "format": "pdf"},
    )


def _parse_docx_legacy(file_bytes: bytes) -> ParsedDocument:
    """Extract text from DOCX using python-docx (paragraphs only, no tables)."""
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    return ParsedDocument(
        text=full_text,
        page_texts=[full_text],
        page_count=1,
        metadata={"parser": "legacy", "format": "docx", "paragraph_count": len(paragraphs)},
    )


def _parse_text_legacy(file_bytes: bytes) -> ParsedDocument:
    """Parse plain text or markdown files."""
    text = file_bytes.decode("utf-8", errors="replace")

    return ParsedDocument(
        text=text,
        page_texts=[text],
        page_count=1,
        metadata={"parser": "legacy", "format": "text"},
    )


# --- Helpers ---

def _get_extension(filename: str) -> str:
    """Extract the file extension (lowercase, including the dot)."""
    dot_idx = filename.rfind(".")
    if dot_idx == -1:
        return ""
    return filename[dot_idx:]
